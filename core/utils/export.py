"""
Утилиты для экспорта истории чата в различные форматы: CSV, DOCX, PDF
"""
import csv
import io
from datetime import datetime
from typing import List, Dict, Any

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from markdown import markdown
    from reportlab.platypus.flowables import Image
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


def export_chat_to_csv(messages: List[Dict[str, Any]], session_title: str = "Chat") -> io.StringIO:
    """
    Экспортирует историю чата в CSV формат.
    
    Args:
        messages: список сообщений с полями role, content, created_at
        session_title: название сессии
    
    Returns:
        StringIO объект с CSV данными
    """
    output = io.StringIO()
    writer = csv.writer(output, quoting=csv.QUOTE_ALL)
    
    # Заголовок
    writer.writerow(['Сессия', session_title])
    writer.writerow(['Дата экспорта', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
    writer.writerow([])
    
    # Заголовки колонок
    writer.writerow(['Дата/Время', 'Роль', 'Сообщение'])
    
    # Сообщения
    for msg in messages:
        created_at = msg.get('created_at', '')
        if isinstance(created_at, datetime):
            created_at = created_at.strftime('%Y-%m-%d %H:%M:%S')
        role = msg.get('role', '')
        content = msg.get('content', '').replace('\n', ' ').replace('\r', '')
        writer.writerow([created_at, role, content])
    
    output.seek(0)
    return output


def export_chat_to_docx(messages: List[Dict[str, Any]], session_title: str = "Chat") -> io.BytesIO:
    """
    Экспортирует историю чата в DOCX формат.
    
    Args:
        messages: список сообщений с полями role, content, created_at
        session_title: название сессии
    
    Returns:
        BytesIO объект с DOCX данными
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    doc = DocxDocument()
    
    # Настройка стилей
    title_style = doc.styles['Heading 1']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading_style = doc.styles['Heading 2']
    heading_style.font.size = Pt(14)
    
    # Заголовок
    title = doc.add_heading(session_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Метаинформация
    meta_para = doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.style.font.size = Pt(10)
    
    doc.add_paragraph()  # Пустая строка
    
    # Сообщения
    for msg in messages:
        created_at = msg.get('created_at', '')
        if isinstance(created_at, datetime):
            created_at = created_at.strftime('%Y-%m-%d %H:%M:%S')
        elif isinstance(created_at, str):
            pass  # Оставляем как есть
        else:
            created_at = str(created_at)
        
        role = msg.get('role', 'unknown')
        content = msg.get('content', '')
        
        # Заголовок сообщения
        role_display = {
            'user': '👤 Пользователь',
            'assistant': '🤖 AI Ассистент',
            'system': '⚙️ Система'
        }.get(role, role)
        
        heading = doc.add_heading(f'{role_display} - {created_at}', level=2)
        
        # Содержимое сообщения (сохраняем markdown форматирование)
        # Простой парсинг markdown в форматированный текст
        content_para = doc.add_paragraph()
        _add_markdown_to_docx(content, content_para)
        
        doc.add_paragraph()  # Пустая строка между сообщениями
    
    # Сохраняем в BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _add_markdown_to_docx(text: str, paragraph):
    """
    Добавляет markdown текст в DOCX параграф с базовым форматированием.
    """
    import re
    
    # Разбиваем на строки
    lines = text.split('\n')
    current_para = paragraph
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Заголовки
        if line.startswith('###'):
            run = current_para.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith('##'):
            run = current_para.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith('#'):
            run = current_para.add_run(line[1:].strip())
            run.bold = True
            run.font.size = Pt(18)
        # Списки
        elif line.startswith('- ') or line.startswith('* '):
            run = current_para.add_run(f'  • {line[2:]}\n')
        elif re.match(r'^\d+\.', line):
            run = current_para.add_run(f'  {line}\n')
        # Жирный текст
        elif '**' in line:
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = current_para.add_run(part[2:-2])
                    run.bold = True
                else:
                    current_para.add_run(part)
        else:
            current_para.add_run(line + '\n')


def export_chat_to_pdf(messages: List[Dict[str, Any]], session_title: str = "Chat") -> io.BytesIO:
    """
    Экспортирует историю чата в PDF формат.
    
    Args:
        messages: список сообщений с полями role, content, created_at
        session_title: название сессии
    
    Returns:
        BytesIO объект с PDF данными
    """
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab не установлен. Установите: pip install reportlab markdown")
    
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    # Стили
    styles = getSampleStyleSheet()
    
    # Кастомные стили
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#3b82f6'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=12,
        leading=14
    )
    
    # Содержимое документа
    story = []
    
    # Заголовок
    title = Paragraph(session_title, title_style)
    story.append(title)
    story.append(Spacer(1, 0.2*inch))
    
    # Метаинформация
    meta = Paragraph(f'Дата экспорта: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal'])
    story.append(meta)
    story.append(Spacer(1, 0.3*inch))
    
    # Сообщения
    for msg in messages:
        created_at = msg.get('created_at', '')
        if isinstance(created_at, datetime):
            created_at = created_at.strftime('%Y-%m-%d %H:%M:%S')
        elif isinstance(created_at, str):
            pass
        else:
            created_at = str(created_at)
        
        role = msg.get('role', 'unknown')
        content = msg.get('content', '')
        
        # Заголовок сообщения
        role_display = {
            'user': '👤 Пользователь',
            'assistant': '🤖 AI Ассистент',
            'system': '⚙️ Система'
        }.get(role, role)
        
        heading_text = f'{role_display} - {created_at}'
        heading = Paragraph(heading_text, heading_style)
        story.append(heading)
        
        # Контент (базовая обработка markdown)
        # Простая замена markdown на HTML для reportlab
        content_html = _markdown_to_html_simple(content)
        para = Paragraph(content_html, normal_style)
        story.append(para)
        
        story.append(Spacer(1, 0.2*inch))
    
    # Создаем PDF
    doc.build(story)
    output.seek(0)
    return output


def _markdown_to_html_simple(text: str) -> str:
    """
    Простое преобразование markdown в HTML для reportlab.
    """
    import re
    
    # Экранируем HTML
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # Заголовки
    text = re.sub(r'^### (.*?)$', r'<b><font size="14">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*?)$', r'<b><font size="16">\1</font></b>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*?)$', r'<b><font size="18">\1</font></b>', text, flags=re.MULTILINE)
    
    # Жирный текст
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    
    # Курсив
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    
    # Списки
    text = re.sub(r'^- (.*?)$', r'  • \1', text, flags=re.MULTILINE)
    
    # Переносы строк
    text = text.replace('\n', '<br/>')
    
    return text

